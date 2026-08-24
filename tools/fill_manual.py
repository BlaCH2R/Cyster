# -*- coding: utf-8 -*-
# Append the missing manual sections (06 properties panel / 07 storyboard
# format / 08 shortcut list) to the Cyster manual copy, following the
# original document's styles: Heading 5 section titles, Normal body, bold
# inline emphasis, and red "❗请注意：" callouts in 黑体.
import docx, io, re
from docx.shared import RGBColor
from docx.oxml.ns import qn

SRC = 'Cyster使用手册(ver.0.1beta) - 副本.docx'

d = docx.Document(SRC)

def set_heiti(run):
    run.font.name = '黑体'
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), '黑体')
    rfonts.set(qn('w:hAnsi'), '黑体')
    rfonts.set(qn('w:eastAsia'), '黑体')

def add_par(text='', bold_segs=None, align=None):
    p = d.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        # **...** 标记为加粗
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                r.bold = True
            else:
                p.add_run(part)
    for seg in (bold_segs or []):
        r = p.add_run(seg[0])
        r.bold = True
    return p

def add_warn(text):
    p = d.add_paragraph()
    p.alignment = 3  # JUSTIFY
    r = p.add_run('❗请注意：')
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    set_heiti(r)
    r2 = p.add_run(text)
    set_heiti(r2)
    return p

def sec(title):
    return d.add_paragraph(title, style='Heading 5')

# ---------------------------------------------------------------------------
sec('06：右侧属性界面介绍')
add_par('右侧的属性界面会跟随当前选中的对象或关键帧实时变化，是Cyster中最主要的编辑入口。'
        '当您在预览画面、对象树或时间轴中选中任意对象时，属性面板会显示该对象的全部可编辑信息；'
        '当播放头不在任何关键帧上时，面板会显示当前时刻的插值结果（只读），便于确认补间动画的实际效果。')
add_par('①对象级字段：属性面板顶部显示对象自身的通用信息——ID（由软件按创建顺序与对象类型自动分配的'
        '唯一识别码）、类型，以及note_controller或携带note选择器对象的note绑定。部分字段'
        '（parent_id、target_id、path、layer、order等）带有黄色“SYNC”标记，它们属于整个时间块的'
        '同步字段，修改后对该对象的全部关键帧同时生效。')
add_par('②状态属性（关键帧字段）：下方是当前选中关键帧的状态属性，包括：时间（支持数字或'
        'start/end/intro/at:NoteID等表达式）、缓动（官方33种缓动完整移植，默认为linear）、'
        '到达后销毁（destroy）；以及场景对象通用字段：坐标x/y/z、旋转rot_x/rot_y/rot_z、'
        '缩放scale/scale_x/scale_y、不透明度opacity、宽度/高度width/height、图层layer、'
        '顺序order、铺满全屏fill_width等。')
add_par('坐标类字段支持直接输入“坐标系:数值”的形式（例如notex:0.5、stagey:300），也可以使用字段'
        '右侧的单位下拉框在stageX/stageY/noteX/noteY/cameraX/cameraY/world之间自动换算。'
        '舞台坐标系为800×600（中心为原点），note坐标系范围为0~1（左下角为原点），相机坐标系基于'
        '相机视图大小，Z轴为3D深度坐标。默认情况下相机位于Z=-10处，向Z轴正方向观察。')
add_par('不同类型的对象会显示各自的特有字段：Sprite为路径path、保持比例preserve_aspect、颜色color；'
        'Text为文本text、字号size、颜色color、对齐align、字间距letter_spacing、字重font_weight'
        '（文本支持b/i/size/color的富文本）；Video为路径path与颜色color；Line为端点列表pos'
        '（每个端点含x/y/z坐标、宽度width与颜色color）。')
add_par('③场景控制器：控制器采用“属性卡片”形式编辑，每个卡片对应一组可独立启用与关键帧化的属性，'
        '包括：不透明度（storyboard不透明度、UI不透明度、扫描线不透明度、背景遮罩、note不透明度'
        '倍率）、扫描线（颜色、覆盖位置）、note颜色（外圈颜色与12项填充色）、相机（透视开关、fov、'
        'x/y/z坐标与旋转），以及全套滤镜卡片（chromatical、bloom、radial_blur、color_adjustment、'
        'color_filter、gray_scale、noise、sepia、dream、fisheye、shockwave、focus、glitch、'
        'arcade、tape）。')
add_par('④Note控制器：用于覆盖谱面中单个或多个note的属性。除note绑定与时间/缓动/销毁外，包含：'
        '覆盖X/Y/Z坐标（override_x/y/z，配合x/y/z、x_multiplier/y_multiplier、dx/dy使用）、'
        '覆盖旋转（override_rot_x/y/z与rot_x/y/z）、覆盖外圈/填充颜色（override_ring_color、'
        'override_fill_color）、不透明度倍率、大小倍率、hitbox倍率、hold方向（hold_direction，'
        '1向上/-1向下，仅Hold）与Hold样式（style，1默认/2下落式）。')
add_warn('官方2.0.2中dy（Y偏移）在下行note（扫描线方向为-1的页面）上存在原生bug，如发现note位置'
         '与预期不符，请在原有值的基础上+1。')
add_par('⑤关键帧操作：属性面板中可以直接在播放头位置“添加关键帧”或“复制当前帧”；时间轴上的关键帧'
        '支持ctrl+c/ctrl+v复制粘贴、按住ctrl多选后统一编辑（数值不一致时显示“多个数值”），以及'
        'Delete键删除。删除仅剩的初始关键帧时会直接删除整个对象；合并时间块（同一轨道内多个对象）'
        '选中后会进入对应的整体属性编辑页。')
add_warn('layer决定对象所在渲染图层——0为默认图层（背景之上、其他游戏元素之下），1在note之下、'
         'UI与背景之上，2在所有游戏元素之上；order决定同一图层内的显示顺序，数值越大越靠上。'
         '请记得为每个sprite设置正确的order，否则可能在正式游戏中出现显示错误'
         '（官方文档建议不确定时设为0）。')

# ---------------------------------------------------------------------------
sec('07：StoryBoard 格式说明')
add_par('本部分依据官方《StoryBoard 格式详解》（v2.0.2）与Cytoid官方仓库源码整理，说明Cyster'
        '编辑与导出的storyboard JSON所对应的格式规格。Cyster的绝大多数编辑操作都会直接生成符合'
        '该格式的JSON，了解这些结构有助于理解时间块与关键帧背后的数据形态。')
add_par('storyboard JSON的顶层包含以下对象组：sprites（图片）、texts（文本）、lines（线段）、'
        'videos（视频，实验性）、controllers（场景控制器）、note_controllers（音符控制器）、'
        'templates（模板）。对象组内的每个对象都包含基准时间time与状态集合states，states中的'
        '每个状态即一个关键帧，软件会在相邻状态之间自动补间。')
add_par('每个对象都至少有一个初始状态（K0）。通用字段包括：id（唯一识别码，支持$note占位符）、'
        'target_id（场景对象专用，使本对象不拥有实体而控制目标对象的实体，只能指向同类对象）、'
        'parent_id（texts/sprites专用，子对象以父对象为坐标系原点并跟随其运动，也可以指向'
        'note_controller实现跟随note）。')
add_par('time为对象的基准时间，支持秒数或以下表达式：“start:NoteID”（note开始时刻）、'
        '“end:NoteID”（结束时刻，Hold/Long Hold用）、“intro:NoteID”（note淡入开始时刻），'
        '三者均可追加“:偏移量”（秒，可为负）；Hold/Long Hold还可使用“at:NoteID:百分比”'
        '（开始时刻+(结束-开始)×百分比，0等同start，1等同end）。未设置time的对象默认不会被启用。')
add_par('relative_time表示相对父状态的时间（父状态时间+relative_time），add_time表示相对最后一个'
        '状态的时间；两者均未定义时使用time。计算顺序：定义了add_time时取“最后状态的时间+add_time”；'
        '定义了relative_time且同时定义了time时取“time+relative_time”；定义了relative_time且有'
        '父状态时取“父状态时间+relative_time”；都未定义时取time。')
add_par('easing为状态间的缓动函数，默认linear；destroy为true时对象完全过渡到该状态后销毁'
        '（官方强烈建议销毁不再需要的对象以提升性能）。states支持嵌套定义，内部states会被拍平追加到'
        '父级状态列表，配合template可一次套用多个状态。')
add_par('场景对象（text/sprite/video/line）的通用状态字段：x/y/z坐标（默认坐标系为'
        'stageX/stageY/深度，z仅在透视相机下有效）、rot_x/rot_y/rot_z旋转、scale/scale_x/'
        'scale_y缩放（scale覆盖前两者）、opacity不透明度（默认0，即所有对象默认不可见）、'
        'width/height尺寸（sprite默认200×200；text自2.0.0起自适应大小）、layer图层（0/1/2）、'
        'order同图层顺序、fill_width铺满屏幕（宽度=屏幕宽、高度=10000，常用于背景替换）。')
add_par('各类对象特有字段：Text——text（支持b/i/size/color富文本）、color（默认#fff）、size'
        '（默认20，动画请用scale而非size）、align（9种对齐，默认middleCenter）、letter_spacing、'
        'font_weight（regular/extraLight/bold/extraBold）；Sprite——path（相对路径，支持'
        '.jpg/.png，建议不超过1920×1080）、preserve_aspect（默认true）、color；Video——path'
        '（建议H.264编码、720p以下mp4）、color（原生游戏中视频在暂停界面不会暂停，为已知问题）；'
        'Line——pos端点数组（每个端点含x/y/z、width默认0.05、color默认#fff），仅支持'
        'opacity/layer/order，常用于绘制扫描线或自定义线段图形。')
add_par('场景控制器控制相机与全局表现：storyboard_opacity（所有storyboard对象透明度）、'
        'ui_opacity（游戏UI透明度）、scanline_opacity（扫描线不透明度）、background_dim'
        '（背景遮罩，默认0.85）、note_opacity_multiplier（note透明度倍率）、scanline_color'
        '（扫描线颜色覆盖）、note_ring_color（note外圈颜色）、note_fill_colors（12项note'
        '填充色：click/drag/hold/long hold/flick/c-drag各上下一组）、override_scanline_pos与'
        'scanline_pos（覆盖扫描线Y坐标）、perspective（透视相机开关，默认true）、size（正交相机'
        '视图大小，默认5）、fov（透视相机视野，默认53.2）、相机x/y/z（默认0/0/-10）与'
        'rot_x/rot_y/rot_z，以及全套滤镜开关与参数。')
add_par('滤镜字段一览：chromatical（+fade/intensity/speed）、bloom（+intensity）、radial_blur'
        '（+intensity，默认0.025）、color_adjustment（brightness/saturation/contrast）、'
        'color_filter（+color）、gray_scale（+intensity）、noise（+intensity，默认0.235）、'
        'sepia（+intensity）、dream（+intensity）、fisheye（+intensity，默认0.5）、'
        'shockwave（+speed）、focus（+size/color/speed/intensity）、glitch（+intensity）、'
        'arcade（+intensity/interference_size/interference_speed/contrast）、tape。')
add_warn('官方文档标注vignette系列与chromatic系列滤镜已在Cytoid 2.0.0中移除，请勿使用；'
         '另外pivot_x/pivot_y、controller的size、text的font等字段在较新引擎中已不生效，'
         'Cyster属性面板不再提供，按旧文档写出这些字段不会产生任何效果。')
add_par('note控制器覆盖谱面note属性：note（目标note id）、override_x/y/z与x/y/z、'
        'x_multiplier/y_multiplier、dx/dy（坐标优先级：设置了override且x有值→x；override且'
        'x为空→原坐标×x_multiplier+dx；否则用原坐标；Y轴同理）、override_rot_x/y/z与'
        'rot_x/y/z（旋转覆盖）、override_ring_color/ring_color与override_fill_color/'
        'fill_color（颜色覆盖）、opacity_multiplier（透明度倍率）、size_multiplier（大小倍率，'
        '官方2.0.2仅对click与flick生效）、hold_direction（hold尾部方向1/-1）、style'
        '（hold样式1/2，2为隐藏连接三角形、尾长渐短、按键特效在hold头播放的下落式样式）。')
add_par('note选择器（note selectors）允许用一个note控制器批量控制符合条件的note，支持按id数组、'
        'type（note类型）、start/end（id范围）、direction（方向）、min_x/max_x（X坐标范围）'
        '筛选，也支持$note占位符。官方文档提示：要让note按曲线运动，请使用两个note控制器分别控制'
        'x与y并设置不同缓动。')
add_warn('正如01章节所述，Cyster默认输出compiled格式storyboard（PascalCase状态、绝对时间、'
         '数值化缓动与Unity颜色结构），这是CytoidPlayer对原生storyboard的一种转换处理，读取速度'
         '与性能更好；因此Cyster内的高级功能（note选择器、$note、相对时间等）在写入文件前会先展开为'
         '绝对形式。请勿在Cyster之外直接修改输出文件或频繁更换谱面，以免往返转换导致读取异常'
         '（详见01章注意事项）。')

# ---------------------------------------------------------------------------
sec('08：快捷键一览')
add_par('以下是Cyster中的常用快捷键，方便快速检索（输入框聚焦时全局快捷键自动让位，避免误触）：')
for line in [
    '空格：播放/暂停',
    '← / →：预览步进0.05秒；Shift+← / Shift+→：步进0.5秒',
    'R：切换预览选择层级（note / stage）',
    'Z：呼出/隐藏预览缩放滑条',
    'Tab：显示/隐藏UI',
    'CapsLock：显示/隐藏所有note',
    'Shift：显示/隐藏note ID',
    'Ctrl+Z / Ctrl+Y：撤销/重做',
    'Ctrl+S：保存项目',
    'Ctrl+C / Ctrl+V：复制/粘贴所选对象或关键帧',
    'Delete / Backspace：删除当前选中',
    'Ctrl+点击：多选对象/关键帧',
]:
    add_par(line)
add_par('以上快捷键大部分也可以在对应界面的按钮提示中看到，如有出入请以按钮提示为准。')

d.save(SRC)

# 简单结构校验：重新打开并统计
d2 = docx.Document(SRC)
added = [p.text for p in d2.paragraphs if p.text.strip().startswith(('06：', '07：', '08：'))]
total = len([p for p in d2.paragraphs if p.text.strip()])
with io.open('tools/fill_manual_report.txt', 'w', encoding='utf-8') as fh:
    fh.write('sections: %s\n' % added)
    fh.write('non-empty paragraphs after: %d\n' % total)
print('saved ok, sections=', added, 'total=', total)
